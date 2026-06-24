from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# 核心配置
IMAGE_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "")
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "output")

STEPS = [
    ("①", "用刀划开原装箱"),
    ("②", "取出一台销售包装"),
    ("③", "用小刀划开封口贴"),
    ("④", "取出旧手册和说明书"),
    ("⑤", "放入新手册和说明书"),
    ("⑥", "还原包装贴上封口贴"),
    ("⑦", "将货物装回箱中"),
]

# 工位分配: {人数: [(工位标签, 步数范围)]}
WORKER_ALLOCATION = {
    1: [("单人", [0, 1, 2, 3, 4, 5, 6])],
    2: [("A", [0, 1, 2, 3]), ("B", [4, 5, 6])],
    3: [("A", [0, 1, 2]), ("B", [3, 4]), ("C", [5, 6])],
    4: [("A", [0, 1]), ("B", [2, 3]), ("C", [4, 5]), ("D", [6])],
}

IMAGES = [
    ("割开封口.jpg", [2]),         # 步骤③
    ("合格证和设备下方的旧产品使用说明书.jpg", [3]),  # 步骤④
    ("设备上方的旧操作说明.jpg", [3]),               # 步骤④
    ("新的操作说明和产品使用说明书.jpg", [4]),       # 步骤⑤
]

QUALITY_NOTES = [
    "不得损坏销售包装",
    "勿遗失三角形合格证",
    "封口贴须贴正、贴牢",
    "新旧手册不得混淆",
]

TOOLS = ["美工刀（含备用刀片）", "封口贴", "新操作手册", "新产品使用说明书", "废料回收箱"]


def set_page_a4_landscape(doc):
    """设置文档为A4横向"""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(0.6)
    section.bottom_margin = Cm(0.6)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)


def add_header(doc, num_workers):
    """添加顶栏标题"""
    title_text = f"增值服务 · 操作手册与说明书置换作业指导卡    {num_workers}人版"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 参数信息行
    info_text = f"客户：中移物流　　地点：广东仓库　　每箱：20台　　版本：{num_workers}人版"
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(info_text)
    run2.font.size = Pt(8)
    run2.font.name = "微软雅黑"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def build_matrix_table(doc, num_workers):
    """构建步骤×工位矩阵表（返回table对象）"""
    allocation = WORKER_ALLOCATION[num_workers]
    worker_labels = [w[0] for w in allocation]

    # 表格: 1列步骤 + len(worker_labels)列工位 = 1+worker_count 列
    # 行: 标题行 + 7步骤行 + 1数量行 = 9行
    num_cols = 1 + len(worker_labels)
    table = doc.add_table(rows=9, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        for i, cell in enumerate(row.cells):
            if i > 0:
                cell.width = Cm(2.0)

    # 标题行
    hdr = table.rows[0]
    hdr.cells[0].text = ""
    run = hdr.cells[0].paragraphs[0].add_run("步骤")
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    for i, label in enumerate(worker_labels):
        cell = hdr.cells[i + 1]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"工位\n{label}")
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 填充颜色函数
    def set_cell_shading(cell, color):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # 数据行 (步骤1-7)
    for step_idx, (step_num, step_desc) in enumerate(STEPS):
        row = table.rows[step_idx + 1]
        # 步骤列
        cell = row.cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(f"{step_num}  {step_desc}")
        run.font.size = Pt(7)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # 工位列
        for w_idx, (_, step_range) in enumerate(allocation):
            cell = row.cells[w_idx + 1]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if step_idx in step_range:
                run = p.add_run("✓")
                run.font.size = Pt(10)
                set_cell_shading(cell, "E8F5E9")  # 浅绿色
            else:
                run = p.add_run("—")
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(200, 200, 200)

    # 末行: 动作数统计
    last_row = table.rows[8]
    cell = last_row.cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("动作数")
    run.bold = True
    run.font.size = Pt(9)

    for w_idx, (_, step_range) in enumerate(allocation):
        cell = last_row.cells[w_idx + 1]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(len(step_range)))
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    return table


def build_image_grid(doc):
    """构建4张参考图网格（返回table对象）"""
    # 创建2行2列的表格用于排列图片
    img_table = doc.add_table(rows=2, cols=2)
    img_table.style = 'Table Grid'
    img_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, (img_filename, step_numbers) in enumerate(IMAGES):
        row_idx = idx // 2
        col_idx = idx % 2
        cell = img_table.cell(row_idx, col_idx)

        img_path = os.path.join(IMAGE_DIR, img_filename)
        step_label = f"[步骤{'/'.join(str(s+1) for s in step_numbers)}]"

        # 添加步骤标签
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(step_label)
        run.font.size = Pt(7)
        run.bold = True
        run.font.color.rgb = RGBColor(0x19, 0x76, 0xD2)

        # 如果是合格证图片，加警告
        if "合格证" in img_filename:
            run2 = p.add_run("  !勿遗失")
            run2.font.size = Pt(7)
            run2.bold = True
            run2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # 添加图片
        if os.path.exists(img_path):
            run_p = cell.add_paragraph()
            run_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_run = run_p.add_run()
            run_run.add_picture(img_path, width=Cm(2.0))
        else:
            print(f"⚠️ 图片文件未找到: {img_path}")
    return img_table


def build_footer_content(cell):
    """在给定的单元格内构建脚注内容（工具清单+质量标准+版本信息）"""
    # 清除单元格默认段落
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)

    run = p.add_run("作业准备")
    run.bold = True
    run.font.size = Pt(8)

    run_q = p.add_run("　　　质量标准")
    run_q.bold = True
    run_q.font.size = Pt(8)

    # 行2: 工具列表 + 质量条目
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)

    tool_text = "  ".join([f"□ {t}" for t in TOOLS])
    quality_text = "  |  ".join([f"• {n}" for n in QUALITY_NOTES])
    line2 = f"{tool_text}　　{quality_text}"
    run2 = p2.add_run(line2)
    run2.font.size = Pt(7)
    run2.font.name = "微软雅黑"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 版本信息行
    p3 = cell.add_paragraph()
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after = Pt(0)
    run3 = p3.add_run("v1.0  编制日期：2026-06-24  编制：科捷物流")
    run3.font.size = Pt(6)
    run3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def generate_sop(num_workers, output_path):
    """生成指定人数版本的SOP"""
    if num_workers not in WORKER_ALLOCATION:
        raise ValueError(f"不支持 {num_workers}人版，仅支持 1/2/3/4")
    doc = Document()
    set_page_a4_landscape(doc)
    add_header(doc, num_workers)

    # 主体区域: 左矩阵 + 右图片 并排，下方脚注
    # 用一个 2 行 2 列的容器表: 第1行=矩阵|图片, 第2行=脚注(合并两列)
    container = doc.add_table(rows=2, cols=2)
    container.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 容器表隐藏边框（仅做布局用）
    tbl = container._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # 设置列宽: 左宽(矩阵) + 右窄(图片)
    for cell in container.columns[0].cells:
        cell.width = Cm(16.0)
    for cell in container.columns[1].cells:
        cell.width = Cm(10.0)

    # 左单元格: 嵌入矩阵表
    left_cell = container.cell(0, 0)
    # 删除单元格默认段落
    left_cell.paragraphs[0].text = ""
    # 构造矩阵表并插入
    matrix_doc = Document()
    m_table = build_matrix_table(matrix_doc, num_workers)
    left_cell._tc.get_or_add_tcPr().append(m_table._tbl)

    # 右单元格: 嵌入图片网格
    right_cell = container.cell(0, 1)
    right_cell.paragraphs[0].text = ""
    img_doc = Document()
    i_table = build_image_grid(img_doc)
    right_cell._tc.get_or_add_tcPr().append(i_table._tbl)

    # 第2行(合并两列): 脚注
    footer_cell = container.cell(1, 0)
    footer_cell.merge(container.cell(1, 1))
    footer_cell.text = ""
    build_footer_content(footer_cell)

    # 设置容器表所有单元格无内边距
    for row in container.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                f'  <w:top w:w="0" w:type="dxa"/>'
                f'  <w:left w:w="30" w:type="dxa"/>'
                f'  <w:bottom w:w="0" w:type="dxa"/>'
                f'  <w:right w:w="30" w:type="dxa"/>'
                f'</w:tcMar>'
            )
            tcPr.append(tcMar)

    doc.save(output_path)
    print(f" SOP文档已生成: {output_path}")


def generate_all_versions():
    """生成1-4人版全部SOP"""
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    for n in [1, 2, 3, 4]:
        output_path = os.path.join(OUTPUT_DIR, f"VAS_SOP_{n}人版.docx")
        generate_sop(n, output_path)
    print(f"\n 全部版本已生成到: {OUTPUT_DIR}")


if __name__ == "__main__":
    generate_all_versions()
